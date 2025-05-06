import io
import base64
import os
import re
import json
from flask import Flask, request, render_template, send_file, redirect, url_for, jsonify, session, flash, send_from_directory
import pandas as pd
import numpy as np
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from google_auth_oauthlib.flow import Flow
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import plotly.express as px      
from docxtpl import DocxTemplate
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
from datetime import datetime
from flask_mail import Mail, Message

app = Flask(__name__)
app.secret_key = 'your_secret_key'
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'

load_dotenv('.env')

UPLOAD_FOLDER = 'uploads'
GENERATED_FOLDER = 'generated'
SETTINGS_PATH = os.path.join(UPLOAD_FOLDER, 'settings.json')
SCOPES = ['https://www.googleapis.com/auth/drive.readonly']
REDIRECT_URI = os.getenv("REDIRECT_URI")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(GENERATED_FOLDER, exist_ok=True)

cached_docx_path = None
cached_csv_path = None
cached_dataframe = None
date_column = None
date_min = None
date_max = None

app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USE_SSL'] = False
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME')
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD')
app.config['MAIL_DEFAULT_SENDER'] = os.getenv('MAIL_USERNAME')

mail = Mail(app)

def generate_chart(df, x_col, y_col, chart_type, output_path, chart_title=None, dpi_scale=2):
    fig = None

    title = chart_title if chart_title else f"{y_col} by {x_col}"

    if chart_type == 'line':
        df[x_col] = pd.to_datetime(df[x_col], errors='coerce')
        daily_df = df.groupby(df[x_col].dt.date).agg({y_col: 'sum'}).reset_index()
        daily_df.rename(columns={daily_df.columns[0]: x_col}, inplace=True)
        fig = px.line(daily_df, x=x_col, y=y_col, title=title, template='simple_white')
    
    elif chart_type == 'bar':
        bar_df = df.groupby(x_col).agg({y_col: 'sum'}).reset_index()
        fig = px.bar(bar_df, x=x_col, y=y_col, title=title, template='simple_white')
    
    elif chart_type == 'hist':
        fig = px.histogram(df, x=y_col, nbins=20, title=title, template='simple_white')
    
    elif chart_type == 'pie':
        pie_data = df.groupby(x_col)[y_col].sum().reset_index()
        fig = px.pie(pie_data, values=y_col, names=x_col, title=title)
    
    else:
        raise ValueError(f"不支援的圖表類型: {chart_type}")

    fig.update_layout(
        template='none',
        font_family=os.path.join("fonts", "cwTeXQYuan-Medium.ttf"),
        title_font_size=20,
        xaxis_title_font_size=16,
        yaxis_title_font_size=16,
        legend_title_font_size=16,
        font_size=14,
        paper_bgcolor='white',
        plot_bgcolor='white'
    )

    fig.write_image(output_path, width=800, height=600, scale=dpi_scale)

def extract_template_variables(template_path):
    doc = DocxTemplate(template_path)
    full_text = ""
    for paragraph in doc.get_docx().paragraphs:
        full_text += paragraph.text + "\n"
    matches = re.findall(r"\{\{\s*(.*?)\s*\}\}", full_text)
    return list(set([match.strip() for match in matches]))

def analyze_csv(csv_path):
    global cached_dataframe, date_column, date_min, date_max

    df = pd.read_csv(csv_path)
    cached_dataframe = df

    for col in df.columns:
        try:
            parsed_dates = pd.to_datetime(df[col], errors='coerce')
            if parsed_dates.notna().sum() > 0:
                date_column = col
                df[col] = parsed_dates
                date_min = parsed_dates.min().date()
                date_max = parsed_dates.max().date()
                break
        except:
            continue

    if not date_column:
        raise Exception("沒有找到日期欄位")

    return df

def convert_docx_to_html(template_path):
    from docx import Document
    document = Document(template_path)
    html = ""

    for para in document.paragraphs:
        line = para.text
        line = re.sub(r"\{\{\s*(.*?)\s*\}\}", r'<span class="variable editable" data-variable="\1">{{\1}}</span>', line)
        html += f"<div>{line}</div>\n"

    return html

def evaluate_formula(formula: str, df: pd.DataFrame, context: dict = None, formulas: dict = None) -> float:
    formula = formula.strip()
    if context is None:
        context = {}
    if formulas is None:
        formulas = {}

    try:
        if formula.startswith("'") and formula.endswith("'"):
            return formula.strip("'")

        col_map = {col.lower(): col for col in df.columns}
        variables_in_formula = re.findall(r"[a-zA-Z_][a-zA-Z0-9_]*", formula)
        
        # 首先處理公式中的變數依賴
        processed_vars = set()
        for var in variables_in_formula:
            if var not in context and var in formulas and var not in processed_vars:
                # 避免循環依賴
                processed_vars.add(var)
                try:
                    # 遞迴計算依賴變數的值
                    next_formula = formulas[var]
                    if isinstance(next_formula, dict) and next_formula.get('type') == 'formula':
                        next_formula = next_formula.get('value', '')
                    context[var] = evaluate_formula(next_formula, df, context, formulas)
                    # print(f"已計算變數 {var} = {context[var]}")
                except Exception as e:
                    # print(f"計算變數 {var} 時發生錯誤: {str(e)}")
                    raise ValueError(f"變數 {var} 計算錯誤: {str(e)}")

        # 處理列名大小寫不敏感
        for var in variables_in_formula:
            var_lower = var.lower()
            if var_lower in col_map:
                correct_col = col_map[var_lower]
                formula = re.sub(r'\b' + re.escape(var) + r'\b', correct_col, formula)

        # 處理特殊函數
        formula = re.sub(r'COUNT\s*\(\s*DISTINCT\s*\((.*?)\)\s*\)', r'df["\1"].nunique()', formula, flags=re.I)
        formula = re.sub(r'MODE\s*\(\s*(.*?)\s*\)', r'df["\1"].mode().iloc[0] if len(df["\1"].mode()) > 0 else None', formula, flags=re.I)
        # 處理特殊的 COUNT(condition) 情況
        count_pattern = re.compile(r'COUNT\s*\(\s*(.+?)\s*\)', re.I)
        for match in count_pattern.finditer(formula):
            content = match.group(1)
            if '==' in content or '!=' in content or '>' in content or '<' in content or '>=' in content or '<=' in content:
                # 條件表達式
                replacement = f'(df.query("{content.replace("==", "==")}", engine="python").shape[0])'
                formula = formula.replace(match.group(0), replacement)
            else:
                # 普通列計數
                replacement = f'df["{content}"].count()'
                formula = formula.replace(match.group(0), replacement)

        formula = re.sub(r'SUM\s*\(\s*(.*?)\s*\)', r'df.eval("\1").sum()', formula, flags=re.I)
        formula = re.sub(r'MEAN\s*\(\s*(.*?)\s*\)', r'df.eval("\1").mean()', formula, flags=re.I)
        formula = re.sub(r'MAX\s*\(\s*(.*?)\s*\)', r'df.eval("\1").max()', formula, flags=re.I)
        formula = re.sub(r'MIN\s*\(\s*(.*?)\s*\)', r'df.eval("\1").min()', formula, flags=re.I)
        formula = re.sub(r'MEDIAN\s*\(\s*(.*?)\s*\)', r'df.eval("\1").median()', formula, flags=re.I)
        formula = re.sub(r'STD\s*\(\s*(.*?)\s*\)', r'df.eval("\1").std()', formula, flags=re.I)
        formula = re.sub(r'VAR\s*\(\s*(.*?)\s*\)', r'df.eval("\1").var()', formula, flags=re.I)

        if formula.upper().startswith("PERCENT_CHANGE("):
            inside = re.findall(r'PERCENT_CHANGE\((.*?)\)', formula, flags=re.I)[0]
            values = df.eval(inside.strip())
            return (values.iloc[-1] - values.iloc[0]) / values.iloc[0] * 100

        if formula.upper().startswith("DIFF("):
            inside = re.findall(r'DIFF\((.*?)\)', formula, flags=re.I)[0]
            values = df.eval(inside.strip())
            return values.iloc[-1] - values.iloc[0]

        if formula.upper().startswith("CAGR("):
            inside = re.findall(r'CAGR\((.*?)\)', formula, flags=re.I)[0]
            parts = [p.strip() for p in inside.split(",")]
            start_val = df.eval(parts[0]).iloc[0]
            end_val = df.eval(parts[1]).iloc[-1]
            periods = float(parts[2])
            return (end_val / start_val) ** (1/periods) - 1

        # 將變數替換為實際值
        for var_name, var_value in context.items():
            if isinstance(var_value, str):
                # 如果是字串，需要加引號
                formula = re.sub(r'\b' + re.escape(var_name) + r'\b', f"'{var_value}'", formula)
            else:
                # 如果是數值，直接替換
                formula = re.sub(r'\b' + re.escape(var_name) + r'\b', str(var_value), formula)

        eval_globals = {"df": df, "np": np}
        eval_globals.update(context)

        result = eval(formula, eval_globals)
        if any(name in formula for name in ["top_region", "top_product"]):
            print(f"⚡ DEBUG: 變數評估後結果 type={type(result)}, value={result}")

        if isinstance(result, (list, pd.Series)):
            if len(result) > 0:
                result = result[0]
            else:
                result = None
        elif isinstance(result, dict):
            if result:
                result = list(result.values())[0]
            else:
                result = None
        return result

    except Exception as e:
        raise ValueError(f"公式錯誤：{str(e)}")

# 新增首頁路由
@app.route('/')
def home():
    message = request.args.get('message', '')
    return render_template('home.html', message=message)

@app.route('/upload', methods=['GET', 'POST'])
def upload():
    global cached_docx_path, cached_csv_path

    if request.method == 'POST':
        if 'docx_file' in request.files:
            docx_file = request.files['docx_file']
            if docx_file.filename != '':
                filename = datetime.now().strftime('%Y%m%d%H%M%S_') + secure_filename(docx_file.filename)
                cached_docx_path = os.path.join(UPLOAD_FOLDER, filename)
                docx_file.save(cached_docx_path)

        if 'csv_file' in request.files:
            csv_file = request.files['csv_file']
            if csv_file.filename != '':
                filename = datetime.now().strftime('%Y%m%d%H%M%S_') + secure_filename(csv_file.filename)
                cached_csv_path = os.path.join(UPLOAD_FOLDER, filename)
                csv_file.save(cached_csv_path)

        if 'settings_file' in request.files:
            settings_file = request.files['settings_file']
            if settings_file.filename != '' and settings_file.filename.endswith('.json'):
                settings_file.save(SETTINGS_PATH)

        if cached_docx_path and cached_csv_path:
            return redirect(url_for('preview'))

    return render_template('index.html')

# 保留原來的 index 路由，但現在轉發到 upload 函數
@app.route('/index', methods=['GET', 'POST'])
def index():
    return upload()

@app.route('/login')
def login():
    flow = Flow.from_client_secrets_file(
        '/etc/secrets/credentials.json',
        scopes=SCOPES,
        redirect_uri='https://reportgen-fn9d.onrender.com/oauth2callback'
    )
    auth_url, state = flow.authorization_url(prompt='consent')
    session['state'] = state
    return redirect(auth_url)

@app.route('/logout')
def logout():
    session.pop('credentials', None)
    return redirect('/')

@app.route('/oauth2callback')
def oauth2callback():
    flow = Flow.from_client_secrets_file(
        '/etc/secrets/credentials.json',
        scopes=SCOPES,
        redirect_uri='https://reportgen-fn9d.onrender.com/oauth2callback'
    )
    flow.fetch_token(authorization_response=request.url)
    creds = flow.credentials
    session['credentials'] = {
        'token': creds.token,
        'refresh_token': creds.refresh_token,
        'token_uri': creds.token_uri,
        'client_id': creds.client_id,
        'client_secret': creds.client_secret,
        'scopes': creds.scopes
    }
    return redirect('/')

@app.route('/set_token', methods=['POST'])
def set_token():
    data = request.get_json()
    token = data.get('token')

    if not token:
        return jsonify({'success': False, 'error': 'Missing token'}), 400

    # 用 token 建立憑證物件
    creds = Credentials(token=token)
    session['credentials'] = {
        'token': creds.token,
        'refresh_token': None,
        'token_uri': "https://oauth2.googleapis.com/token",
        'client_id': os.getenv("GOOGLE_CLIENT_ID"),
        'client_secret': '',  # 可以留空，不需要設定
        'scopes': SCOPES
    }

    return jsonify({'success': True})


@app.route('/save_token', methods=['POST'])
def save_token():
    data = request.get_json()
    session['google_access_token'] = data.get('access_token')
    return jsonify({"success": True})

@app.route('/preview', methods=['GET'])
def preview():
    global cached_docx_path, cached_csv_path

    if not cached_docx_path or not cached_csv_path:
        return redirect(url_for('upload'))

    try:
        variables = extract_template_variables(cached_docx_path)
        analyze_csv(cached_csv_path)
        html_content = convert_docx_to_html(cached_docx_path)

        return render_template('preview.html',
                            html_content=html_content,
                            variables=variables,
                            date_column=date_column,
                            date_min=date_min,
                            date_max=date_max)
    except Exception as e:
        print(f"預覽生成錯誤: {str(e)}")
        return render_template('index.html', error=f"預覽生成失敗: {str(e)}")

@app.route('/filter_data', methods=['POST'])
def filter_data():
    global cached_dataframe, date_column

    data = request.json
    start_date = pd.to_datetime(data.get('start_date'))
    end_date = pd.to_datetime(data.get('end_date'))

    if not date_column:
        return jsonify([])

    filtered = cached_dataframe[
        (cached_dataframe[date_column] >= start_date) &
        (cached_dataframe[date_column] <= end_date)
    ]

    return jsonify(filtered.to_dict(orient='records'))

@app.route('/render_preview', methods=['POST'])
def render_preview():
    data = request.json.get('data', [])
    formulas = request.json.get('formulas', {})

    if not data:
        return jsonify({})

    df = pd.DataFrame(data)
    results = {}
    calculated_context = {}
    special_vars = ['start_date', 'end_date']

    # 建立公式依賴關係圖
    formula_graph = {}
    for var, formula_info in formulas.items():
        if isinstance(formula_info, dict) and formula_info.get('type') == 'formula':
            formula = formula_info.get('value', '')
        elif isinstance(formula_info, dict) and formula_info.get('type') == 'chart':
            formula = ''
        else:
            formula = formula_info

        dependencies = set()
        for other_var in formulas:
            if other_var != var and isinstance(formula, str) and re.search(r'\b' + re.escape(other_var) + r'\b', formula):
                dependencies.add(other_var)
        formula_graph[var] = dependencies

    visited = set()
    temp_visited = set()
    order = []

    def topo_sort(node):
        if node in temp_visited:
            print(f"循環依賴: {node}")
            return
        if node in visited:
            return
        temp_visited.add(node)
        for dep in formula_graph.get(node, []):
            topo_sort(dep)
        temp_visited.remove(node)
        visited.add(node)
        order.append(node)

    for var in formulas:
        if var not in visited:
            topo_sort(var)

    order.reverse()

    for var in formulas:
        if var not in order:
            try:
                setting = formulas[var]
                if var in special_vars and 'date' in df.columns:
                    if var == 'start_date':
                        value = df['date'].min()
                    else:
                        value = df['date'].max()
                    if pd.isna(value):
                        value = ''
                    else:
                        value = pd.to_datetime(value).strftime('%Y/%m/%d')
                    results[var] = value
                    calculated_context[var] = value
                elif isinstance(setting, dict):
                    if setting.get('type') == 'formula':
                        expr = setting.get('value')
                        value = evaluate_formula(expr, df, context=calculated_context, formulas=formulas)
                        if hasattr(value, "item"):
                            value = value.item()
                        if isinstance(value, (int, float)):
                            value = round(value, 2)
                        results[var] = value
                        calculated_context[var] = value
                    elif setting.get('type') == 'chart':
                        results[var] = ''
                    else:
                        results[var] = setting.get('value', '')
                        calculated_context[var] = setting.get('value', '')
                else:
                    results[var] = setting
                    calculated_context[var] = setting
            except Exception as e:
                results[var] = f"錯誤: {str(e)}"

    for var in order:
        try:
            setting = formulas[var]
            if isinstance(setting, dict):
                if setting.get('type') == 'chart':
                    x_col = setting.get('x')
                    y_col = setting.get('y')
                    chart_type = setting.get('chartType')

                    if not all([x_col, y_col, chart_type]):
                        results[var] = '錯誤：缺少圖表設定'
                        continue

                    fig = None
                    if chart_type == 'line':
                        df[x_col] = pd.to_datetime(df[x_col], errors='coerce')
                        daily_df = df.groupby(df[x_col].dt.date).agg({y_col: 'sum'}).reset_index()
                        daily_df.rename(columns={daily_df.columns[0]: x_col}, inplace=True)
                        fig = px.line(daily_df, x=x_col, y=y_col, title=f"{y_col} by {x_col}")
                    
                    elif chart_type == 'bar':
                        bar_df = df.groupby(x_col).agg({y_col: 'sum'}).reset_index()
                        fig = px.bar(bar_df, x=x_col, y=y_col, title=f"{y_col} by {x_col}")
                    
                    elif chart_type == 'hist':
                        fig = px.histogram(df, x=y_col, nbins=20, title=f"{y_col} Histogram")
                    
                    elif chart_type == 'pie':
                        pie_data = df.groupby(x_col)[y_col].sum().reset_index()
                        fig = px.pie(pie_data, values=y_col, names=x_col, title=f"{y_col} 分佈")
                    
                    else:
                        results[var] = '錯誤：不支援的圖表類型'
                        continue

                    save_path = os.path.join('generated', f'{var}.png')
                    fig.write_image(save_path, width=800, height=600, scale=2)

                    results[var] = ''

                elif setting.get('type') == 'formula':
                    expr = setting.get('value')
                    value = evaluate_formula(expr, df, context=calculated_context, formulas=formulas)
                    if hasattr(value, "item"):
                        value = value.item()
                    if isinstance(value, (int, float)):
                        value = round(value, 2)
                    results[var] = value
                    calculated_context[var] = value

                else:
                    results[var] = setting.get('value', '')
                    calculated_context[var] = setting.get('value', '')
            else:
                results[var] = setting
                calculated_context[var] = setting

        except Exception as e:
            results[var] = f"錯誤: {str(e)}"

    return jsonify(results)

@app.route('/render', methods=['POST'])
def render_word():
    global cached_docx_path # Ensure this global variable is accessible

    formulas = request.json.get('formulas', {})
    data = request.json.get('data', [])
    filename = request.json.get('filename', 'final_report.docx')

    if not cached_docx_path or not os.path.exists(cached_docx_path):
         return "錯誤：找不到 Word 模板文件。", 400
    if not data:
        return "錯誤：沒有提供用於渲染的數據。", 400

    try:
        filtered_df = pd.DataFrame(data)
        # Convert date column to datetime if it exists and isn't already
        # Assuming date_column is known globally or can be inferred
        # if date_column and date_column in filtered_df.columns:
        #     filtered_df[date_column] = pd.to_datetime(filtered_df[date_column], errors='coerce')

    except Exception as e:
        print(f"Error creating DataFrame from data: {e}")
        return f"錯誤：處理輸入數據時出錯: {str(e)}", 400

    context = {}
    calculated_context = {} # Keep track of calculated values for dependencies

    # --- Explicitly add start_date and end_date ---
    # Ensure they are set before potentially complex formula evaluations.
    if 'start_date' in formulas and isinstance(formulas['start_date'], dict) and formulas['start_date'].get('type') == 'fixed':
        start_date_val = formulas['start_date'].get('value', '')
        context['start_date'] = start_date_val
        calculated_context['start_date'] = start_date_val
        print(f"DEBUG (render_word): Explicitly set start_date in context: {start_date_val}")

    if 'end_date' in formulas and isinstance(formulas['end_date'], dict) and formulas['end_date'].get('type') == 'fixed':
        end_date_val = formulas['end_date'].get('value', '')
        context['end_date'] = end_date_val
        calculated_context['end_date'] = end_date_val
        print(f"DEBUG (render_word): Explicitly set end_date in context: {end_date_val}")

    # --- Topological Sort Logic ---
    formula_graph = {}
    valid_vars = set(formulas.keys()) # Keep track of variables defined in formulas

    for var, formula_info in formulas.items():
        # Skip dates if already handled
        if var in ['start_date', 'end_date'] and var in context:
             continue

        formula = '' # Default to empty string
        if isinstance(formula_info, dict):
            if formula_info.get('type') == 'formula':
                formula = formula_info.get('value', '')
        elif isinstance(formula_info, str): # Handle case where formula is just a string
             formula = formula_info
        # else: # Charts or other types have no formula string for dependency check

        dependencies = set()
        if isinstance(formula, str) and formula: # Only check dependencies if formula is a non-empty string
            # Find potential variable names (simple regex, might need refinement)
            potential_deps = re.findall(r'[a-zA-Z_][a-zA-Z0-9_]*', formula)
            for dep in potential_deps:
                # Check if the potential dependency is a defined variable, not the variable itself,
                # and not a known function name (basic check, might need expansion)
                 if dep in valid_vars and dep != var and dep.upper() not in ['SUM', 'MEAN', 'MAX', 'MIN', 'COUNT', 'MEDIAN', 'STD', 'VAR', 'PERCENT_CHANGE', 'DIFF', 'CAGR', 'MODE', 'DISTINCT', 'NP']:
                     dependencies.add(dep)
        formula_graph[var] = dependencies


    visited = set()
    recursion_stack = set() # To detect cycles
    order = []

    def topo_sort(node):
        if node not in valid_vars: # Skip if node isn't a defined variable
             print(f"DEBUG (topo_sort): Node '{node}' not in valid_vars, skipping.")
             return
        # Skip dates if already handled
        if node in ['start_date', 'end_date'] and node in context:
            return

        if node in recursion_stack:
            print(f"錯誤：偵測到循環依賴 involving '{node}'")
            raise Exception(f"循環依賴: {node}") # Raise error on cycle
        if node in visited:
            return

        visited.add(node)
        recursion_stack.add(node)

        # Ensure dependencies exist in the graph before recursing
        for dep in formula_graph.get(node, []):
             if dep in valid_vars:
                 topo_sort(dep)
             else:
                 # This dependency might be a column name, not another formula variable.
                 # Or it could be an error in the formula. Let evaluate_formula handle it later.
                 print(f"DEBUG (topo_sort): Dependency '{dep}' for node '{node}' not in formula graph keys.")


        recursion_stack.remove(node)
        # Only add nodes that are actually defined in formulas to the order
        order.append(node) # Add after visiting all dependencies

    # Build the execution order
    for var in formulas:
         if var not in visited and var not in ['start_date', 'end_date']: # Avoid re-processing handled dates
            try:
                topo_sort(var)
            except Exception as e:
                 print(f"Error during topological sort for var '{var}': {e}")
                 # Return error if cycle detected during sort
                 return f"公式計算順序錯誤（可能存在循環依賴）: {str(e)}", 500


    # No need to reverse if appending happens after visiting dependencies
    print(f"DEBUG (render_word): Formula calculation order: {order}")

    # --- Initialize DocxTemplate ---
    try:
        doc = DocxTemplate(cached_docx_path)
    except Exception as e:
         print(f"Error loading DocxTemplate: {e}")
         return f"錯誤：無法加載 Word 模板 '{os.path.basename(cached_docx_path)}': {str(e)}", 500


    # --- Process variables ---
    # Process variables NOT necessarily in the determined order first,
    # especially fixed values or those without dependencies that weren't explicitly handled (like dates).
    for var, setting in formulas.items():
        if var not in order and var not in context: # Check if not already processed by explicit handling or order
            print(f"DEBUG (render_word): Processing non-ordered var: {var}")
            try:
                value_to_set = None
                is_chart = False
                if isinstance(setting, dict):
                    var_type = setting.get('type')
                    if var_type == 'formula':
                         # Evaluate only if no dependencies or handled dependencies
                        expr = setting.get('value', '')
                        # Be cautious evaluating here, dependencies might not be ready.
                        # Better to rely on the 'order' loop for formulas.
                        # Maybe set as placeholder? For now, let's attempt evaluation.
                        value_to_set = evaluate_formula(expr, filtered_df, context=calculated_context, formulas=formulas)
                    elif var_type == 'fixed':
                         value_to_set = setting.get('value', '')
                    elif var_type == 'chart':
                         is_chart = True # Mark as chart, handle in the main loop
                         context[var] = "[圖表將在此生成]" # Placeholder
                    else: # Unknown dict type
                         value_to_set = str(setting) # Convert dict to string as fallback

                else: # Handle cases where 'setting' is not a dictionary (e.g., simple value)
                    # Could be a simple formula string - attempt evaluation
                    try:
                         value_to_set = evaluate_formula(str(setting), filtered_df, context=calculated_context, formulas=formulas)
                    except ValueError: # If evaluation fails, treat as literal string
                         value_to_set = str(setting)


                # Common post-processing for non-chart numeric values
                if not is_chart and value_to_set is not None:
                    if hasattr(value_to_set, "item"): # Convert numpy types
                        value_to_set = value_to_set.item()
                    if isinstance(value_to_set, (int, float)):
                         value_to_set = round(value_to_set, 2) # Round numeric results

                    context[var] = value_to_set
                    calculated_context[var] = value_to_set # Update calculated context
                    print(f"DEBUG (render_word): Set non-ordered var '{var}' in context: {value_to_set}")

            except Exception as e:
                print(f"錯誤：處理非順序變數 '{var}' 時出錯: {e}")
                context[var] = f"[錯誤: {str(e)}]"
                calculated_context[var] = None # Mark as error in calculated context


    # --- Process variables IN topological order ---
    for var in order:
        # Ensure the variable is actually in formulas before processing
        if var not in formulas:
             print(f"警告 (render_word): 變數 '{var}' 在計算順序中但不在 formulas 中，跳過。")
             continue
        # Skip if already processed (e.g., fixed values handled above)
        if var in context and not (isinstance(formulas[var], dict) and formulas[var].get('type') == 'chart' and context[var] == "[圖表將在此生成]"):
            print(f"DEBUG (render_word): Var '{var}' already in context, skipping ordered processing.")
            continue

        print(f"DEBUG (render_word): Processing ordered var: {var}")
        try:
            setting = formulas[var]
            value_to_set = None # Reset for each variable
            is_chart = False

            if isinstance(setting, dict):
                var_type = setting.get('type')
                if var_type == 'formula':
                    expr = setting.get('value', '')
                    print(f"DEBUG (render_word): Evaluating formula for '{var}': {expr}")
                    value_to_set = evaluate_formula(expr, filtered_df, context=calculated_context, formulas=formulas)

                elif var_type == 'chart':
                    is_chart = True
                    x_col = setting.get('x')
                    y_col = setting.get('y')
                    chart_type = setting.get('chartType')
                    chart_title = setting.get('chartTitle') # Get chart title
                    # Use a default DPI scale if not provided
                    dpi_scale_str = setting.get('dpi', '2') # Get DPI setting as string
                    try:
                        dpi_scale = float(dpi_scale_str)
                    except (ValueError, TypeError):
                        print(f"警告：圖表 '{var}' 的 DPI 設定 '{dpi_scale_str}' 無效，使用預設值 2。")
                        dpi_scale = 2


                    if not x_col or not y_col or not chart_type:
                        print(f"警告 (render_word): 圖表 '{var}' 缺少設定 (X:{x_col}, Y:{y_col}, Type:{chart_type})，跳過。")
                        context[var] = "[圖表設定不完整]"
                        continue
                     # Ensure columns exist in DataFrame
                    if x_col not in filtered_df.columns or y_col not in filtered_df.columns:
                         print(f"錯誤 (render_word): 圖表 '{var}' 所需欄位 ({x_col}, {y_col}) 不在資料中，跳過。")
                         context[var] = f"[錯誤：找不到欄位 {x_col} 或 {y_col}]"
                         continue

                    img_path = os.path.join(GENERATED_FOLDER, f"{var}.png")
                    print(f"DEBUG (render_word): Generating chart for '{var}' with title: '{chart_title}', DPI scale: {dpi_scale}")
                    generate_chart(
                        filtered_df,
                        x_col,
                        y_col,
                        chart_type,
                        img_path, # Pass output path directly
                        chart_title=chart_title, # Pass the title
                        dpi_scale=dpi_scale # Pass DPI scale
                        )
                    # Add InlineImage to context if image generation succeeded
                    if os.path.exists(img_path):
                         context[var] = InlineImage(doc, img_path, width=Mm(120)) # Adjust width as needed
                         print(f"DEBUG (render_word): Added chart '{var}' to context.")
                    else:
                         print(f"錯誤 (render_word): 圖表 '{var}' 圖片生成失敗或未找到。")
                         context[var] = "[圖表生成失敗]"
                    # Charts don't usually add to calculated_context unless their result is needed elsewhere

                elif var_type == 'fixed': # Should have been handled above, but catch just in case
                    value_to_set = setting.get('value', '')
                else: # Unknown dict type
                     value_to_set = str(setting)

            else: # Handle non-dict formulas/values if they end up in 'order'
                 print(f"DEBUG (render_word): Evaluating non-dict formula/value for '{var}': {setting}")
                 try:
                     value_to_set = evaluate_formula(str(setting), filtered_df, context=calculated_context, formulas=formulas)
                 except ValueError as ve:
                     # If evaluation fails, treat as literal string? Or report error?
                     print(f"警告：評估非字典變數 '{var}' 時出錯 ({ve})，將其視為字串。")
                     value_to_set = str(setting)

            # Common post-processing and context update for non-chart variables
            if not is_chart and value_to_set is not None:
                 if hasattr(value_to_set, "item"): # Convert numpy types
                     value_to_set = value_to_set.item()
                 if isinstance(value_to_set, (int, float)):
                    # Apply rounding only if it's not already a string representation
                    if not isinstance(value_to_set, str):
                         value_to_set = round(value_to_set, 2)

                 context[var] = value_to_set
                 calculated_context[var] = value_to_set # Update calculated context
                 print(f"DEBUG (render_word): Set ordered var '{var}' in context: {value_to_set}")
            elif not is_chart and value_to_set is None:
                 print(f"警告 (render_word): 變數 '{var}' 計算結果為 None。")
                 context[var] = "" # Set empty string for None results in template? Or handle differently?
                 calculated_context[var] = None


        except Exception as e:
            print(f"錯誤：處理順序變數 '{var}' 時發生嚴重錯誤: {e}")
            import traceback
            traceback.print_exc() # Print full traceback for debugging
            context[var] = f"[處理錯誤: {str(e)}]"
            calculated_context[var] = None # Mark as error in calculated context

    # --- Final rendering ---
    try:
        print(f"DEBUG (render_word): Final context for rendering: {context}") # Debug: print final context
        doc.render(context)
        output_path = os.path.join(GENERATED_FOLDER, filename)
        # Ensure the generated folder exists
        os.makedirs(GENERATED_FOLDER, exist_ok=True)
        doc.save(output_path)
        print(f"INFO (render_word): Word document saved to {output_path}")
        return send_file(output_path, as_attachment=True, download_name=filename)
    except Exception as e:
         print(f"錯誤：最終渲染 Word 文件 '{filename}' 時出錯: {e}")
         import traceback
         traceback.print_exc()
         # Consider returning an error response to the user
         return f"渲染 Word 時發生嚴重錯誤: {str(e)}", 500

@app.route('/save_settings', methods=['POST'])
def save_settings():
    settings = request.json
    formulas = settings.get('formulas', {})

    # 自動正規化 formulas
    new_formulas = {}
    for var, setting in formulas.items():
        if isinstance(setting, dict) and 'type' in setting:
            new_formulas[var] = setting
        else:
            # 舊格式，補成 type=formula
            new_formulas[var] = {
                "type": "formula",
                "value": setting
            }

    # 重新組成完整 settings
    new_settings = {
        "formulas": new_formulas
    }

    with open(SETTINGS_PATH, 'w', encoding='utf-8') as f:
        json.dump(new_settings, f, ensure_ascii=False, indent=2)
    return "設定已儲存", 200

@app.route('/load_settings', methods=['GET'])
def load_settings():
    if not os.path.exists(SETTINGS_PATH):
        return jsonify({}), 404

    with open(SETTINGS_PATH, 'r', encoding='utf-8') as f:
        settings = json.load(f)

    formulas = settings.get('formulas', {})
    new_formulas = {}
    for var, setting in formulas.items():
        if isinstance(setting, dict) and 'type' in setting:
            new_formulas[var] = setting
        else:
            # 舊格式，補成 type=formula
            new_formulas[var] = {
                "type": "formula",
                "value": setting
            }

    settings['formulas'] = new_formulas
    return jsonify(settings)

@app.route('/get_columns', methods=['GET'])
def get_columns():
    global cached_dataframe
    if cached_dataframe is not None:
        return jsonify({"columns": list(cached_dataframe.columns)})
    else:
        return jsonify({"columns": []})

@app.route('/generated/<path:filename>')
def serve_generated_file(filename):
    return send_from_directory('generated', filename)

@app.route('/regenerate_chart', methods=['POST'])
def regenerate_chart():
    content = request.get_json()

    var_name = content.get('varName')
    x_col = content.get('x')
    y_col = content.get('y')
    chart_type = content.get('chartType')
    chart_title = content.get('chartTitle')
    data = content.get('data')
    dpi_scale = content.get('dpi', 2)  # 🔥 預設 dpi_scale 2

    if not all([var_name, x_col, y_col, chart_type, data]):
        return jsonify({'error': '缺少必要參數'}), 400

    df = pd.DataFrame(data)

    if df.empty:
        return jsonify({'error': '沒有資料'}), 400

    try:
        save_path = os.path.join('generated', f'{var_name}.png')
        generate_chart(df, x_col, y_col, chart_type, save_path, chart_title=chart_title, dpi_scale=dpi_scale)  # 🔥 使用上面的 generate_chart
        return jsonify({'message': '重新產生完成'})
    except Exception as e:
        return jsonify({'error': str(e)})

@app.route('/get_chart_preview', methods=['POST'])
def get_chart_preview():
    content = request.get_json()
    var_name = content.get('varName')  # 🔥 傳進來變數名
    
    if not var_name:
        return jsonify({'error': '缺少變數名稱'}), 400

    filename = f"{var_name}.png"  # 圖片檔案命名規則
    filepath = os.path.join('generated', filename)

    if not os.path.exists(filepath):
        return jsonify({'error': '找不到圖片'}), 404

    # 🔥 回傳圖片 URL，而不是 base64
    return jsonify({'image_url': url_for('static', filename=f'../generated/{filename}')})


# 聯絡表單處理路由
@app.route('/contact', methods=['POST'])
def contact():
    name = request.form.get('name')
    email = request.form.get('email')
    message = request.form.get('message')

    if not name or not email or not message:
        flash('請完整填寫所有欄位喔～', 'warning')
        return redirect(url_for('home'))

    try:
        msg = Message('網站聯絡表單新訊息 ✉️', recipients=[app.config['MAIL_USERNAME']], charset='utf-8')
        msg.body = f"""
收到一封新的聯絡表單：

姓名：{name}
電子郵件：{email}
訊息內容：
{message}
        """
        mail.send(msg)
        flash('感謝您的聯絡！我們已收到您的訊息！', 'success')
    except Exception as e:
        print(f"寄信失敗：{e}")
        flash('寄信時發生錯誤，請稍後再試！', 'warning')

    return redirect(url_for('home'))

# 處理示例模板下載的路由
@app.route('/download_sample', methods=['GET'])
def download_sample():
    sample_path = os.path.join('static', 'samples', 'sample_template.docx')
    # 確保目錄存在
    os.makedirs(os.path.dirname(sample_path), exist_ok=True)
    
    # 如果文件不存在，創建一個簡單的樣本文件
    if not os.path.exists(sample_path):
        from docx import Document
        doc = Document()
        doc.add_heading('範例模板', 0)
        doc.add_paragraph('這是一個範例模板，您可以在其中使用 {{ 變數名稱 }} 格式來定義需要替換的變數。')
        p = doc.add_paragraph('例如：本月銷售總額為 ')
        p.add_run('{{ 銷售總額 }}').bold = True
        p.add_run('，較上月')
        p.add_run('{{ 銷售增長率 }}').bold = True
        doc.add_paragraph('系統將自動計算並填充這些變數的值。')
        doc.save(sample_path)
    
    return send_file(sample_path, as_attachment=True)

@app.route('/config')
def get_config():
    return jsonify({
        "GOOGLE_CLIENT_ID": os.getenv("GOOGLE_CLIENT_ID"),
        "GOOGLE_DEVELOPER_KEY": os.getenv("GOOGLE_DEVELOPER_KEY")
    })

@app.route('/import_drive_file')
def import_drive_file():
    global cached_docx_path, cached_csv_path

    file_id = request.args.get('file_id')
    file_type = request.args.get('type')  # 可用來判斷是哪種檔案：docx、csv、json

    if not file_id or not file_type:
        return jsonify(success=False, error="Missing parameters")

    if 'credentials' not in session:
        return jsonify(success=False, error="Not logged in to Google")

    try:
        creds = Credentials(**session['credentials'])
        drive_service = build('drive', 'v3', credentials=creds)
        metadata = drive_service.files().get(fileId=file_id).execute()
        file_name = metadata['name']
        mime_type = metadata['mimeType']

        ext = {
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/csv': 'csv',
            'application/json': 'json'
        }.get(mime_type)

        if not ext:
            return jsonify(success=False, error="Unsupported file type")

        save_folder = 'uploads'
        os.makedirs(save_folder, exist_ok=True)
        timestamped_name = datetime.now().strftime('%Y%m%d%H%M%S_') + secure_filename(file_name)
        save_path = os.path.join(save_folder, timestamped_name)

        request_drive = drive_service.files().get_media(fileId=file_id)
        with io.FileIO(save_path, 'wb') as fh:
            downloader = MediaIoBaseDownload(fh, request_drive)
            done = False
            while not done:
                status, done = downloader.next_chunk()

        # 🔥 根據檔案類型設為 cached_xxx_path
        if ext == 'docx':
            cached_docx_path = save_path
        elif ext == 'csv':
            cached_csv_path = save_path
        # 設定檔不用設 cached，但可擴充處理
        return jsonify(success=True, filename=timestamped_name, file_type=ext, mime_type=mime_type)

    except Exception as e:
        return jsonify(success=False, error=str(e))

if __name__ == '__main__':
    app.run(debug=True, port="8000")